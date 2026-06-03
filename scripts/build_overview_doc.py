"""
scripts/build_overview_doc.py — Generate the JD Builder client briefing document.

Run once to produce docs/JD_Builder_System_Overview.docx.
Not a template — this is a standalone briefing document for internal classification clients.

Usage:
    python scripts/build_overview_doc.py
"""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

OUTPUT_PATH = "docs/JD_Builder_System_Overview.docx"

# GC Federal Identity palette
GC_RED = RGBColor(0xEB, 0x00, 0x0B)
GC_NAVY = RGBColor(0x26, 0x37, 0x4A)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
MID_GREY = RGBColor(0x99, 0x99, 0x99)


def _set_para_spacing(para, before: int = 0, after: int = 6) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _shade_row(row, fill_hex: str) -> None:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)


def _cell(cell, text: str, *, bold: bool = False, italic: bool = False,
          color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ------------------------------------------------------------------ #
    # Header bar — red rule + department line                             #
    # ------------------------------------------------------------------ #
    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = dept.add_run("Department of National Defence  |  ADM(HR-Civ)  |  Classification")
    r.font.size = Pt(9)
    r.font.color.rgb = MID_GREY
    r.bold = False
    _set_para_spacing(dept, before=0, after=4)

    # Red divider
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(0)
    divider.paragraph_format.space_before = Pt(0)
    pPr = divider._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "EB000B")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ------------------------------------------------------------------ #
    # Title block                                                          #
    # ------------------------------------------------------------------ #
    doc.add_paragraph()  # spacer

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("JD Builder — System Overview")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = GC_NAVY
    _set_para_spacing(title, before=0, after=4)

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run(
        f"How the tool works and where your feedback matters    ·    {date.today().strftime('%B %d, %Y')}"
    )
    sr.font.size = Pt(10)
    sr.italic = True
    sr.font.color.rgb = MID_GREY
    _set_para_spacing(subtitle, before=0, after=200)

    # ------------------------------------------------------------------ #
    # Section helper                                                       #
    # ------------------------------------------------------------------ #
    def section_heading(text: str) -> None:
        h = doc.add_paragraph()
        r = h.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = GC_RED
        r.font.all_caps = True
        _set_para_spacing(h, before=180, after=40)
        # Underline via border
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "EB000B")
        pBdr.append(bot)
        pPr.append(pBdr)

    def body(text: str, *, italic: bool = False, before: int = 0, after: int = 80) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.italic = italic
        _set_para_spacing(p, before=before, after=after)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        r = p.runs[0] if p.runs else p.add_run("")
        # Replace content
        for run in p.runs:
            run.text = ""
        run = p.add_run(text)
        run.font.size = Pt(10)
        _set_para_spacing(p, before=0, after=40)

    # ------------------------------------------------------------------ #
    # What it is                                                           #
    # ------------------------------------------------------------------ #
    section_heading("What It Is")
    body(
        "JD Builder is a guided web tool that helps classification advisors draft a complete, "
        "defensible Work Description (WD) for a DND position — step by step. At the end, it "
        "exports a populated Word document."
    )

    # ------------------------------------------------------------------ #
    # The Workflow                                                          #
    # ------------------------------------------------------------------ #
    section_heading("The Workflow")
    body("The tool walks an advisor through five sequential steps:", after=60)

    steps = [
        ("1  Describe the position",
         "The advisor pastes or types a plain-language description of the role."),
        ("2  NOC mapping",
         "The tool searches the National Occupational Classification (NOC 2021) database and proposes "
         "candidate NOC codes, ranked by relevance. The advisor reviews and confirms one."),
        ("3  Occupational Group (OG) classification",
         "Using the confirmed NOC code and TBS OG definitions, the tool proposes the occupational group "
         "(e.g. AS, EC, PM). The advisor confirms or overrides."),
        ("4  Draft duties",
         "The AI selects the most relevant Main Duties statements from the NOC element library for the "
         "confirmed NOC code, filtered to fit the confirmed OG. Critically: the tool only picks from "
         "existing, authoritative NOC duty statements — it never writes new text. Advisors can add "
         "their own duty statements on top. The tool also runs an \"orphan check,\" flagging any duty "
         "that may fall outside the OG's functional scope based on the official OG inclusion/exclusion rules."),
        ("5  JES scoring",
         "For each Job Evaluation Standard (JES) factor (e.g. skills, effort, responsibility, working "
         "conditions), the AI reviews the duties and proposes a degree level and point value. Each factor "
         "is scored separately and the advisor can retry individual factors or override with their own "
         "judgment and rationale."),
        ("6  Export",
         "Once all JES factors are complete, the tool generates a .docx Work Description. For DND "
         "positions, it also allows linking the role to Departmental Results Framework (DRF) outputs "
         "before export."),
    ]

    tbl = doc.add_table(rows=len(steps), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Inches(1.6), Inches(4.4)]
    for i, (step_label, step_desc) in enumerate(steps):
        row = tbl.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        _cell(row.cells[0], step_label, bold=True, color=GC_NAVY, size=10)
        _cell(row.cells[1], step_desc, size=10)
        if i % 2 == 0:
            _shade_row(row, "EFF3F8")

    doc.add_paragraph()

    # ------------------------------------------------------------------ #
    # Key Design Principles                                                #
    # ------------------------------------------------------------------ #
    section_heading("Key Design Principles")

    principles = [
        ("The AI never writes the JD.",
         "Duty statements come verbatim from the NOC database. The AI's only job is to select and rank "
         "from pre-approved source material."),
        ("Every statement is traceable.",
         "Each duty carries a provenance tag: which NOC code, which version of the NOC, and when it "
         "was retrieved. Advisor-added statements are tagged as ADVISOR."),
        ("Advisor override at every step.",
         "NOC, OG, individual duties, and any JES factor can all be manually overridden — with rationale "
         "recorded — so the advisor always has final say."),
        ("Orphan detection.",
         "The tool checks whether proposed duties actually fall within the OG's defined scope, helping "
         "catch misclassification before a WD is finalized."),
    ]

    for label, desc in principles:
        p = doc.add_paragraph()
        r1 = p.add_run(label + "  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = GC_NAVY
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)
        _set_para_spacing(p, before=60, after=60)

    # ------------------------------------------------------------------ #
    # Data Sources                                                         #
    # ------------------------------------------------------------------ #
    section_heading("Data Sources Currently in Use")

    sources = [
        ("NOC 2021 (StatCan)", "Occupational codes and Main Duties statement banks"),
        ("TBS OG Definitions", "Occupational group boundaries (inclusions, exclusions)"),
        ("CA-JES (TBS)", "Factor definitions, degree descriptors, and point values"),
        ("DND DRF Dataset", "Departmental results and core responsibilities for linkage"),
        ("TBS OG Pay Rates", "Reference data for rates of pay"),
    ]

    src_tbl = doc.add_table(rows=len(sources) + 1, cols=2)
    src_tbl.style = "Table Grid"
    _cell(src_tbl.rows[0].cells[0], "Source", bold=True, color=GC_NAVY)
    _cell(src_tbl.rows[0].cells[1], "What It Provides", bold=True, color=GC_NAVY)
    _shade_row(src_tbl.rows[0], "26374A")
    # Fix header text color (cell() sets font color but shading overwrites visually — ensure white)
    for cell in src_tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (source, desc) in enumerate(sources):
        row = src_tbl.rows[i + 1]
        _cell(row.cells[0], source, bold=True, size=10)
        _cell(row.cells[1], desc, size=10)
        if i % 2 == 0:
            _shade_row(row, "EFF3F8")

    doc.add_paragraph()

    # ------------------------------------------------------------------ #
    # What it is NOT yet                                                   #
    # ------------------------------------------------------------------ #
    section_heading("What It Is Not (Yet)")

    not_yet = [
        "Does not consult OCHRO classification rulings or past WD precedents",
        "Does not handle collective agreement nuances or dual-occupation positions",
        "Does not validate against the organizational hierarchy or DND establishment data",
        "DRF linkage step is new and lightly tested",
    ]
    for item in not_yet:
        bullet(item)

    doc.add_paragraph()

    # ------------------------------------------------------------------ #
    # Where feedback matters                                               #
    # ------------------------------------------------------------------ #
    section_heading("Where Your Feedback Matters Most")
    body(
        "We built this tool with classification accuracy as the primary constraint. Your "
        "expert critique at this stage directly shapes what gets built next. Specifically, "
        "we are looking for input on:"
    )

    feedback_items = [
        "Are the right data sources driving each step?",
        "Are there classification scenarios (dual OG, atypical duties, senior positions, "
        "AS/EC distinction) where the proposed logic would clearly break?",
        "Are there guardrails that are missing — or guardrails that would get in the way of "
        "legitimate advisor judgment?",
        "Is the orphan-statement check useful in practice, or does it generate noise?",
        "Does the JES scoring proposal give you something useful to react to, or is it "
        "too far off to be a meaningful starting point?",
    ]
    for item in feedback_items:
        bullet(item)

    # ------------------------------------------------------------------ #
    # Footer                                                               #
    # ------------------------------------------------------------------ #
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    fr = footer_p.add_run(
        f"DRAFT — INTERNAL USE ONLY    ·    Prepared {date.today().strftime('%B %d, %Y')}    ·    "
        "ADM(HR-Civ) Classification Modernization"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = MID_GREY
    fr.italic = True
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
