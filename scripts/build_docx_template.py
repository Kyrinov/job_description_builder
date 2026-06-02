"""
scripts/build_docx_template.py — Generate the docxtpl TBS Work Description template.

Run once to (re)generate templates/docx/work_description_template.docx. The .docx
is a committed binary artifact — the contract that the export service (Plan 08-02)
fills via a context dict. The Jinja2 variable names listed below are stable;
renaming them is a contract break.

Template structure (TBS Work Description format, D-04):
    Title:          "Work Description"
    Section 1:      Position Identification (table: label/value rows)
    Section 2:      Organizational Context (paragraph + italic source citation)
    Section 3:      Key Activities (Duties) — paragraph-level {%p for %} loop
    Section 4:      Classification — JES Scoring — table-row {%tr for %} loop
    Section 5:      Source Document Version Manifest — table-row {%tr for %} loop

Jinja2 variables (contract):
    position_title, position_number, og_level, supervisor_title,
    supervisor_position_number, review_date,
    organizational_context_text, organizational_context_source,
    duties (list of {text, source_id, source_version, is_advisor}),
    jes_scores (list of {factor_name, level, points, source_id, source_version}),
    jes_total_points,
    manifest (list of {source_type, source_id, source_version, retrieved_date})

Re-run this script to update the template; then commit the regenerated .docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate

OUTPUT_PATH = "templates/docx/work_description_template.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Replace the cell's content with a single paragraph containing the given text."""
    # Clear any default empty paragraph that python-docx places in the cell
    for para in list(cell.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)
    para = cell.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def build() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    title = doc.add_heading("Work Description", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Section 1: Position Identification
    # ------------------------------------------------------------------
    doc.add_heading("1. Position Identification", level=1)

    position_rows = [
        ("Position Title:", "{{ position_title }}"),
        ("Position Number:", "{{ position_number }}"),
        ("Group and Level:", "{{ og_level }}"),
        ("Supervisor Title:", "{{ supervisor_title }}"),
        ("Supervisor Position Number:", "{{ supervisor_position_number }}"),
        ("Review Date:", "{{ review_date }}"),
    ]
    pos_table = doc.add_table(rows=len(position_rows), cols=2)
    pos_table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(position_rows):
        _set_cell_text(pos_table.rows[i].cells[0], label, bold=True)
        _set_cell_text(pos_table.rows[i].cells[1], value)

    # ------------------------------------------------------------------
    # Section 2: Organizational Context
    # ------------------------------------------------------------------
    doc.add_heading("2. Organizational Context", level=1)
    doc.add_paragraph("{{ organizational_context_text }}")
    src_para = doc.add_paragraph()
    src_run = src_para.add_run("Source: {{ organizational_context_source }}")
    src_run.italic = True

    # ------------------------------------------------------------------
    # Section 3: Key Activities (Duties) — paragraph-level loop
    # ------------------------------------------------------------------
    doc.add_heading("3. Key Activities (Duties)", level=1)

    # Each {%p ... %} / {%p endfor %} must be in its own paragraph (per docxtpl:
    # "Do not use {%p twice in the same paragraph").
    doc.add_paragraph("{%p for duty in duties %}")
    doc.add_paragraph("{{ duty.text }}")
    src_para = doc.add_paragraph()
    src_para.add_run("Source: {{ duty.source_id }} ({{ duty.source_version }})").italic = True
    # Advisor marker (D-06): conditional paragraph with literal label.
    # if/endif each in their own paragraph; the body line between them is
    # rendered only when duty.is_advisor is True.
    doc.add_paragraph("{%p if duty.is_advisor %}")
    doc.add_paragraph("[advisor-added / not from authoritative source]")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endfor %}")

    # ------------------------------------------------------------------
    # Section 4: Classification — JES Scoring — table-row loop
    # ------------------------------------------------------------------
    doc.add_heading("4. Classification — JES Scoring", level=1)

    # 4 rows: header, for-marker, data, endfor-marker.
    # The for and endfor MUST be in their own rows (each with only the tag in
    # the first cell) — docxtpl's patch_xml regex is greedy and matches the
    # LAST {%tr %} tag in a row, so co-locating {%tr for %} and {%tr endfor %}
    # in the same data row would cause the for tag to be eaten.
    jes_table = doc.add_table(rows=4, cols=4)
    jes_table.style = "Light Grid Accent 1"
    _set_cell_text(jes_table.rows[0].cells[0], "Factor", bold=True)
    _set_cell_text(jes_table.rows[0].cells[1], "Level", bold=True)
    _set_cell_text(jes_table.rows[0].cells[2], "Points", bold=True)
    _set_cell_text(jes_table.rows[0].cells[3], "Source", bold=True)
    # Row 1: {%tr for %} marker
    _set_cell_text(jes_table.rows[1].cells[0], "{%tr for f in jes_scores %}")
    # Row 2: data row (duplicated per factor by the for loop)
    _set_cell_text(jes_table.rows[2].cells[0], "{{ f.factor_name }}")
    _set_cell_text(jes_table.rows[2].cells[1], "{{ f.level }}")
    _set_cell_text(jes_table.rows[2].cells[2], "{{ f.points }}")
    _set_cell_text(jes_table.rows[2].cells[3], "{{ f.source_id }} ({{ f.source_version }})")
    # Row 3: {%tr endfor %} marker
    _set_cell_text(jes_table.rows[3].cells[0], "{%tr endfor %}")

    # Total Points line
    total_para = doc.add_paragraph()
    total_run = total_para.add_run("Total Points: ")
    total_run.bold = True
    total_para.add_run("{{ jes_total_points }}")

    # ------------------------------------------------------------------
    # Section 5: Source Document Version Manifest — table-row loop
    # ------------------------------------------------------------------
    doc.add_heading("5. Source Document Version Manifest", level=1)

    # Same for/data/endfor pattern as the JES table.
    manifest_table = doc.add_table(rows=4, cols=4)
    manifest_table.style = "Light Grid Accent 1"
    _set_cell_text(manifest_table.rows[0].cells[0], "Source", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[1], "ID", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[2], "Version", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[3], "Date", bold=True)
    _set_cell_text(manifest_table.rows[1].cells[0], "{%tr for m in manifest %}")
    _set_cell_text(manifest_table.rows[2].cells[0], "{{ m.source_type }}")
    _set_cell_text(manifest_table.rows[2].cells[1], "{{ m.source_id }}")
    _set_cell_text(manifest_table.rows[2].cells[2], "{{ m.source_version }}")
    _set_cell_text(manifest_table.rows[2].cells[3], "{{ m.retrieved_date }}")
    _set_cell_text(manifest_table.rows[3].cells[0], "{%tr endfor %}")

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")

    # Self-verify: load the generated template with docxtpl and list variables.
    # This catches malformed template tags at build time, not at first export.
    tpl = DocxTemplate(OUTPUT_PATH)
    undeclared = sorted(tpl.get_undeclared_template_variables())
    print(f"Template variables ({len(undeclared)}): {undeclared}")


if __name__ == "__main__":
    build()
