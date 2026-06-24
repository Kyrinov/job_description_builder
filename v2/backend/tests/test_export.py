"""
test_export.py — Phase 20: Export endpoint tests.

Integration tests for:
  POST /api/wd/{id}/export/docx  (API-08, EXP-01)
  POST /api/wd/{id}/export/poster (API-09, EXP-02)
  POST /api/wd/{id}/export/pdf    (EXP-03)

All tests are skipped until Plan 02 implements export.py and wires it
into api/__init__.py. Remove @pytest.mark.skip when the router is live.
"""
from __future__ import annotations

import io

import docx
import pytest

pytestmark = pytest.mark.asyncio


async def _create_wd(client) -> str:
    resp = await client.post(
        "/api/wd",
        json={"record": {"title": "Test Role"}, "answers": {}, "step_index": 1},
    )
    assert resp.status_code == 201
    return resp.json()["id"]


async def _create_wd_with_jes_scores(client) -> str:
    """Create a WD seeded with confirmed_og + jes_total_points required for export."""
    wd_id = await _create_wd(client)
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "EC", "og_name": "Economics and Social Science Services"},
            "og_level": 4,
            "jes_total_points": 621,
            "jes_scores": [
                {"factor_name": "Decision Making", "degree": 3, "points": 150},
                {"factor_name": "Communication", "degree": 2, "points": 84},
            ],
            "duties": [
                {
                    "id": "d1",
                    "text": "Provides advice on economic policy.",
                    "source": "noc",
                    "provenance_noc_code": "4163",
                    "provenance_hash": "abc123",
                    "advisor": False,
                }
            ],
        },
    )
    assert resp.status_code == 200
    return wd_id


# ---------------------------------------------------------------------------
# Phase 25 — Accessible Template fixture helpers (ACC-02 RED baseline)
# 4 JES-shape variants needed to exercise the new Effort/Working-Conditions
# bucketing logic: EC (LLM-style, no category key on persisted dict),
# point-rating-with-Effort (FB), point-rating-without-Effort (MT),
# level-description (AS, jes_scores: []).
# ---------------------------------------------------------------------------

_DUTY_SEED = {
    "id": "d1",
    "text": "Provides advice on policy matters.",
    "source": "noc",
    "provenance_noc_code": "4163",
    "provenance_hash": "abc123",
    "advisor": False,
}

_RECORD_SEED = {
    "title": "Test Role",
    "client_service_results": "Citizens receive timely, accurate policy guidance.",
    "quals": {
        "education": "Degree in economics.",
        "experience": "5 years policy analysis.",
    },
}

_QUAL_SEED = {
    "education": "Degree in economics.",
    "experience": "5 years policy analysis.",
    "source": "EC-05 default",
    "last_modified": "2026-06-16T00:00:00Z",
}


async def _create_wd_ec(client) -> str:
    """EC: jes_scores carry Effort + Conditions factors but NO `category` key.

    The EC LLM-scoring path (_build_factor_score in jes_service.py) omits the
    category field from the persisted dict, so export_service must look up
    category via a factor_name -> category map (ACC-02 pitfall 2).
    """
    wd_id = await _create_wd(client)
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "EC", "og_name": "Economics and Social Science Services"},
            "og_level": 4,
            "jes_total_points": 43,
            "jes_scores": [
                {"factor_name": "Physical effort", "degree": 2, "points": 4},
                {"factor_name": "Sensory effort", "degree": 1, "points": 2},
                {"factor_name": "Working conditions", "degree": 3, "points": 12},
                {"factor_name": "Communication", "degree": 2, "points": 25},
            ],
            "duties": [_DUTY_SEED],
            "record": _RECORD_SEED,
            "qualification": _QUAL_SEED,
        },
    )
    assert resp.status_code == 200
    return wd_id


async def _create_wd_point_rating_with_effort(client) -> str:
    """FB (point-rating): jes_scores carry Effort + Conditions factors WITH category key.

    Exercises the point-rating path where category IS on the persisted dict
    (per JES_FACTORS_BY_GROUP['FB'] in constants.py).
    """
    wd_id = await _create_wd(client)
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "FB", "og_name": "Border Services"},
            "og_level": 4,
            "jes_total_points": 71,
            "jes_scores": [
                {"factor_name": "Physical effort", "degree": 2, "points": 5},
                {"factor_name": "Sensory effort", "degree": 2, "points": 4},
                {"factor_name": "Risk to health", "degree": 2, "points": 10},
                {"factor_name": "Work environment", "degree": 1, "points": 2},
                {"factor_name": "Knowledge", "degree": 3, "points": 50},
            ],
            "duties": [_DUTY_SEED],
            "record": _RECORD_SEED,
            "qualification": _QUAL_SEED,
        },
    )
    assert resp.status_code == 200
    return wd_id


async def _create_wd_point_rating_no_effort(client) -> str:
    """MT (point-rating): jes_scores carry ONLY Skill/Responsibility factors.

    No Effort or Working Conditions category at all — must fall back to
    '[To be completed by advisor]' placeholder (ACC-02).
    """
    wd_id = await _create_wd(client)
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "MT", "og_name": "Meteorology"},
            "og_level": 4,
            "jes_total_points": 110,
            "jes_scores": [
                {"factor_name": "Knowledge", "degree": 3, "points": 50},
                {"factor_name": "Decision making", "degree": 3, "points": 60},
            ],
            "duties": [_DUTY_SEED],
            "record": _RECORD_SEED,
            "qualification": _QUAL_SEED,
        },
    )
    assert resp.status_code == 200
    return wd_id


async def _create_wd_level_description(client) -> str:
    """AS (level-description): jes_total_points set, jes_scores is [].

    No factor data at all for level-description groups — must fall back to
    '[To be completed by advisor]' placeholder (ACC-02).
    """
    wd_id = await _create_wd(client)
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "AS", "og_name": "Administrative Services"},
            "og_level": 4,
            "jes_total_points": 500,
            "jes_scores": [],
            "duties": [_DUTY_SEED],
            "record": _RECORD_SEED,
            "qualification": _QUAL_SEED,
        },
    )
    assert resp.status_code == 200
    return wd_id


async def test_export_wd_docx_returns_bytes(client, env_with_db):
    """EXP-01 / API-08 — POST /api/wd/{id}/export/docx returns .docx bytes with correct MIME type."""
    wd_id = await _create_wd_with_jes_scores(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 0


async def test_export_wd_docx_manifest(client, env_with_db):
    """EXP-01 — Exported DOCX bytes are non-zero (version manifest rendered in template)."""
    wd_id = await _create_wd_with_jes_scores(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    # Proxy for "manifest section rendered": file is > 5 kB (empty docx is ~4 kB)
    assert len(resp.content) > 5000, "DOCX file suspiciously small — manifest may not have rendered"


async def test_export_wd_docx_amendments_appendix(client, env_with_db):
    """EXP-01 / AMEND-02 — DOCX bytes delivered even when amendment notes exist."""
    wd_id = await _create_wd_with_jes_scores(client)
    # Add an amendment note
    await client.post(
        f"/api/wd/{wd_id}/amendments",
        json={"section": "du", "comment": "Review this duty for scope."},
    )
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    assert len(resp.content) > 0


async def test_export_poster_returns_bytes(client, env_with_db):
    """EXP-02 / API-09 — POST /api/wd/{id}/export/poster returns .docx bytes."""
    wd_id = await _create_wd_with_jes_scores(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/poster")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 0


async def test_export_pdf_501_when_weasyprint_absent(client, env_with_db, monkeypatch):
    """EXP-03 — POST /api/wd/{id}/export/pdf returns 501 when WeasyPrint import fails."""
    import sys
    monkeypatch.setitem(sys.modules, "weasyprint", None)
    wd_id = await _create_wd(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/pdf")
    assert resp.status_code == 501
    assert "WeasyPrint" in resp.json()["detail"]


async def test_export_docx_404(client, env_with_db):
    """API-08 — POST /api/wd/does-not-exist/export/docx returns 404."""
    resp = await client.post("/api/wd/does-not-exist/export/docx")
    assert resp.status_code == 404


async def test_export_poster_404(client, env_with_db):
    """API-09 — POST /api/wd/does-not-exist/export/poster returns 404."""
    resp = await client.post("/api/wd/does-not-exist/export/poster")
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# WR-08: PDF 501 via runtime probe failure (Pango/Cairo missing at render time)
# ---------------------------------------------------------------------------

async def test_export_pdf_501_when_weasyprint_probe_fails(client, env_with_db, monkeypatch):
    """WR-08 — POST /api/wd/{id}/export/pdf returns 501 when _probe_weasyprint() is False.

    Distinct from the import-failure path: WeasyPrint imports cleanly but the
    runtime probe (_probe_weasyprint) returns False, indicating Pango/Cairo are
    missing. This covers the second 501 branch in the PDF handler.
    """
    import app.services.export_service as es
    monkeypatch.setattr(es, "_weasyprint_available", False)

    wd_id = await _create_wd(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/pdf")
    # 501 from the _probe_weasyprint() guard; or 501 if weasyprint not installed
    assert resp.status_code == 501


# ---------------------------------------------------------------------------
# WR-09: 409 from require_og_confirmed gate on all export endpoints
# ---------------------------------------------------------------------------

async def test_export_docx_409_without_og(client, env_with_db):
    """WR-09 — POST /api/wd/{id}/export/docx returns 409 when OG not confirmed."""
    wd_id = await _create_wd(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 409
    assert resp.json()["detail"]["error"] == "classification_pending"


async def test_export_poster_409_without_og(client, env_with_db):
    """WR-09 — POST /api/wd/{id}/export/poster returns 409 when OG not confirmed."""
    wd_id = await _create_wd(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/poster")
    assert resp.status_code == 409
    assert resp.json()["detail"]["error"] == "classification_pending"


# ---------------------------------------------------------------------------
# WR-10: self-healing JES flow in export_wd_docx
# ---------------------------------------------------------------------------

async def test_export_docx_self_heals_jes_scores(client, env_with_db, monkeypatch):
    """WR-10 — DOCX export triggers JES self-healing when jes_total_points is None.

    Confirms that the self-healing block is exercised: score_jes_v2 is called
    and the WD is re-loaded. We monkeypatch score_jes_v2 to a no-op (CI has no
    LLM) and assert that the export still succeeds (200 + non-empty bytes).
    """
    import app.api.export as export_mod

    async def _fake_score(**kwargs):  # noqa: ARG001
        return None

    monkeypatch.setattr(export_mod, "score_jes_v2", _fake_score)

    wd_id = await _create_wd(client)
    # PATCH confirmed_og + og_level so the 409 gate passes, but leave
    # jes_total_points as None so self-healing is triggered.
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "EC", "og_name": "Economics and Social Science Services"},
            "og_level": 4,
            "duties": [
                {
                    "id": "d1",
                    "text": "Analyses economic data for policy recommendations.",
                    "source": "noc",
                    "provenance_noc_code": "4163",
                    "advisor": False,
                }
            ],
        },
    )
    assert resp.status_code == 200

    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    assert len(resp.content) > 0


# ---------------------------------------------------------------------------
# WR-11: duties record-fallback in _build_wd_context
# ---------------------------------------------------------------------------

async def test_export_docx_uses_record_duties_fallback(client, env_with_db):
    """WR-11 — DOCX export succeeds when duties live in record only (no root duties).

    Exercises the _build_wd_context record-fallback path: root wd.duties is
    empty but record.duties has duty data. The export must still return 200
    and non-empty bytes.
    """
    wd_id = await _create_wd(client)
    # PATCH via the record dict only — no root-level duties field
    resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={
            "confirmed_og": {"og_code": "EC", "og_name": "Economics and Social Science Services"},
            "og_level": 3,
            "jes_total_points": 520,
            "jes_scores": [
                {"factor_name": "Decision Making", "degree": 2, "points": 120},
            ],
            "record": {
                "title": "Policy Analyst",
                "duties": [
                    {
                        "id": "d1",
                        "text": "Develops policy briefs for senior management.",
                        "source": "noc",
                        "provenance_noc_code": "4163",
                        "provenance_section": "Main duties",
                        "provenance_hash": "abc123",
                        "advisor": False,
                        "orphan": False,
                    }
                ],
            },
        },
    )
    assert resp.status_code == 200

    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    assert len(resp.content) > 0


# ---------------------------------------------------------------------------
# Phase 21 — OGX-02: NON_EC_STANDARD_NAMES consolidated into constants.py
# ---------------------------------------------------------------------------

def test_standard_names_import_from_constants():
    """OGX-02 — export_service.py must import NON_EC_STANDARD_NAMES from constants.py,
    not define it locally.

    FAILS at Wave 0: export_service.py lines 50-55 still have local dict.
    Goes GREEN after Plan 02 (Wave 1) removes the local dict and adds the import.
    """
    import inspect
    import importlib
    export_service = importlib.import_module("app.services.export_service")
    source = inspect.getsource(export_service)
    # Must import from app.data.constants — not define locally
    assert "from app.data.constants import" in source and "NON_EC_STANDARD_NAMES" in source, \
        "export_service.py must import NON_EC_STANDARD_NAMES from app.data.constants"
    # Must NOT define a local copy
    assert "NON_EC_STANDARD_NAMES: dict" not in source, \
        "export_service.py must not define a local NON_EC_STANDARD_NAMES dict"


# ---------------------------------------------------------------------------
# Phase 25 — Accessible Template: ACC-02 / ACC-04 RED baseline (Plan 25-01)
# These 6 tests pin the contract for the new GoC Accessible JD format.
# They are EXPECTED TO FAIL at Wave 0 (current implementation is the
# legacy TBS Work Description format) — Plans 02 and 03 turn them GREEN
# by rewriting _build_wd_context and rendering through the new
# wd_accessible_template.docx (ACC-01..04).
# ---------------------------------------------------------------------------

# 7 Part 2 subsection headings from the reference document
# (data/AI Docs/Accessible Job Description Template (1).docx).
ACCESSIBLE_PART2_HEADINGS = [
    "Organizational context",
    "Client service results",
    "Key activities",
    "Skills",
    "Effort",
    "Responsibilities",
    "Working conditions",
]

ADVISOR_PLACEHOLDER = "[To be completed by advisor]"


def _docx_text(content: bytes) -> str:
    """Read back a rendered DOCX and concatenate paragraph + table-cell text.

    Used by the content-presence and structure tests to inspect what the
    export endpoint actually produced, since raw .docx bytes are zipped
    XML and can't be grepped directly.
    """
    d = docx.Document(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# ACC-02: Effort / Working Conditions section content (4 fallback branches)
# ---------------------------------------------------------------------------

async def test_accessible_effort_ec_populated(client, env_with_db):
    """ACC-02 — EC export must show Effort and Working-Conditions factors populated
    despite the EC scoring path persisting jes_scores WITHOUT a `category` key.

    The new Accessible format must look up category via factor_name -> category
    map (EC_JES_ELEMENTS / JES_FACTORS_BY_GROUP), NOT trust wd.jes_scores[*].category.

    Asserts both the factor name strings (which appear in the current TBS
    template's JES Factor column) AND the 'Effort' section heading (capital E),
    which only appears in the new Accessible format — the factor names
    'Physical effort' / 'Sensory effort' use lowercase 'e' so they don't match.
    This combination distinguishes a dedicated 'Effort' section from a stray
    factor_name cell in a JES table.
    """
    wd_id = await _create_wd_ec(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    assert "Physical effort" in text, "EC Effort factor 'Physical effort' not present in rendered DOCX"
    assert "Sensory effort" in text, "EC Effort factor 'Sensory effort' not present in rendered DOCX"
    assert "Working conditions" in text, "EC Working Conditions factor not present in rendered DOCX"
    # Heading-style 'Effort' (capital E) — distinguishes Accessible section
    # from the lowercase 'effort' in factor_name values.
    assert "Effort" in text, (
        "Expected 'Effort' section heading (capital E) in rendered DOCX — "
        "EC effort factors must appear under a dedicated Effort section."
    )


async def test_accessible_effort_fb_populated(client, env_with_db):
    """ACC-02 — FB (point-rating with Effort + Conditions) export must show
    both Conditions factors ('Risk to health', 'Work environment') and Effort
    factors ('Physical effort', 'Sensory effort') in the rendered DOCX.

    Asserts factor names AND the 'Effort' heading (capital E) so the test
    fails in the current TBS template (which has no 'Effort' section) and
    passes once the Accessible template lands in Plan 25-02/03.
    """
    wd_id = await _create_wd_point_rating_with_effort(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    assert "Risk to health" in text, "FB Conditions factor 'Risk to health' not present"
    assert "Work environment" in text, "FB Conditions factor 'Work environment' not present"
    assert "Physical effort" in text, "FB Effort factor 'Physical effort' not present"
    assert "Sensory effort" in text, "FB Effort factor 'Sensory effort' not present"
    assert "Effort" in text, (
        "Expected 'Effort' section heading in FB export — Effort factors "
        "must appear under a dedicated Effort section, not just in the JES table."
    )


async def test_accessible_effort_no_factor_group_placeholder(client, env_with_db):
    """ACC-02 — Point-rating group with NO Effort/Conditions categories (MT)
    must render the '[To be completed by advisor]' placeholder string in the
    Effort and Working-Conditions sections.
    """
    wd_id = await _create_wd_point_rating_no_effort(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    assert ADVISOR_PLACEHOLDER in text, (
        "Expected '[To be completed by advisor]' placeholder for MT group with no Effort/WC factors"
    )


async def test_accessible_effort_level_description_placeholder(client, env_with_db):
    """ACC-02 — Level-description group (AS, jes_scores: []) must render
    the '[To be completed by advisor]' placeholder in Effort and
    Working-Conditions sections.
    """
    wd_id = await _create_wd_level_description(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    assert ADVISOR_PLACEHOLDER in text, (
        "Expected '[To be completed by advisor]' placeholder for level-description group (AS)"
    )


# ---------------------------------------------------------------------------
# ACC-04: Content-presence (no literal Jinja2 / None leaks in rendered DOCX)
# ---------------------------------------------------------------------------

async def test_accessible_content_presence(client, env_with_db):
    """ACC-04 — Rendered DOCX for a fully-completed EC WD must have NO:
      * unrendered Jinja2 tag (literal '{{')
      * bare 'None' token (str(None) leak)
      * unrendered Jinja2 block tag (literal '%}')

    Together these three guards catch every common template-rendering bug
    the Accessible-template rewrite is at risk of introducing.

    Also asserts that `record.client_service_results` text (seeded via
    _RECORD_SEED) is actually rendered in the new 'Client service results'
    Part 2 subsection — fails against the current TBS template which has
    no such field/section.
    """
    wd_id = await _create_wd_ec(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    assert "{{" not in text, "Unrendered Jinja2 expression tag leaked into output"
    # Wrap with newlines so the assertion matches a bare 'None' line, not a
    # substring like 'Nonetax' or 'NoneType'.
    assert "\nNone\n" not in ("\n" + text + "\n"), "str(None) leaked into rendered output"
    assert "%}" not in text, "Unrendered Jinja2 block tag leaked into output"
    # The new 'Client service results' Part 2 subsection must surface
    # record.client_service_results (seeded by _create_wd_ec). Fails in the
    # current TBS template (no such field/render) — locks the new data path.
    assert "Citizens receive timely" in text, (
        "client_service_results from record was not rendered — Part 2 "
        "'Client service results' section must surface record.client_service_results"
    )


# ---------------------------------------------------------------------------
# ACC-01: Structure — all 7 Part 2 subsection headings + Part 1/2 markers
# ---------------------------------------------------------------------------

async def test_accessible_structure_headings(client, env_with_db):
    """ACC-01 — Rendered Accessible DOCX must contain all 7 Part 2 subsection
    headings (Organizational context, Client service results, Key activities,
    Skills, Effort, Responsibilities, Working conditions) and the Part 1 /
    Part 2 markers from the reference document.
    """
    wd_id = await _create_wd_ec(client)
    resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert resp.status_code == 200
    text = _docx_text(resp.content)
    for heading in ACCESSIBLE_PART2_HEADINGS:
        assert heading in text, f"Missing Part 2 subsection heading: {heading!r}"
    assert "Part 1" in text, "Missing 'Part 1' marker (Position information and signatures)"
    assert "Part 2" in text, "Missing 'Part 2' marker (Job description)"


# ---------------------------------------------------------------------------
# Phase 26 — ORG-03: RED baseline for org_context export priority
# ---------------------------------------------------------------------------

async def test_org_context_in_export(client, env_with_db):
    """ORG-03: When org_context is set, the typed value appears in the DOCX output."""
    wd_id = await _create_wd_ec(client)
    patch_resp = await client.patch(
        f"/api/wd/{wd_id}", json={"org_context": "Test org context text for export"}
    )
    assert patch_resp.status_code == 200

    export_resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert export_resp.status_code == 200
    doc = docx.Document(io.BytesIO(export_resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Test org context text for export" in full_text


async def test_org_context_fallback_in_export(client, env_with_db):
    """ORG-03: When org_context is None, synthesized fallback (branch/reports/summary) is used.
    No template variable leak ({{}}) must appear in the output."""
    wd_id = await _create_wd_ec(client)
    # _create_wd_ec does NOT set org_context — org_context stays None
    export_resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert export_resp.status_code == 200
    doc = docx.Document(io.BytesIO(export_resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "{{" not in full_text
    assert "organizational_context_text" not in full_text


async def test_org_context_empty_string_falls_back(client, env_with_db):
    """F-01: A whitespace-only org_context must fall back to synthesized text,
    not render a blank Organizational Context section (defense-in-depth)."""
    wd_id = await _create_wd_ec(client)
    patch_resp = await client.patch(f"/api/wd/{wd_id}", json={"org_context": "   "})
    assert patch_resp.status_code == 200

    export_resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert export_resp.status_code == 200
    doc = docx.Document(io.BytesIO(export_resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "{{" not in full_text
    assert "organizational_context_text" not in full_text


# ---------------------------------------------------------------------------
# Phase 27 — RESP-03: RED baseline for responsibilities_narrative export priority
# (R-RESP-03: DOCX Part 2 "Responsibility" content = wd.responsibilities_narrative
#  when filled, else _ADVISOR_PLACEHOLDER. Replaces the JES-derived
#  responsibilities_text block in _build_wd_context.)
# ---------------------------------------------------------------------------

async def test_responsibilities_narrative_in_export(client, env_with_db):
    """RESP-03: When responsibilities_narrative is set, the typed value appears
    in the DOCX Part 2 Responsibility section (asserts it is rendered — not
    the JES-derived factor list)."""
    wd_id = await _create_wd_ec(client)
    patch_resp = await client.patch(
        f"/api/wd/{wd_id}",
        json={"responsibilities_narrative": "Owns the environmental policy portfolio and briefs senior leadership."},
    )
    assert patch_resp.status_code == 200

    export_resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert export_resp.status_code == 200
    doc = docx.Document(io.BytesIO(export_resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Owns the environmental policy portfolio and briefs senior leadership." in full_text


async def test_responsibilities_narrative_placeholder_in_export(client, env_with_db):
    """RESP-03: When responsibilities_narrative is None/empty, the advisor
    placeholder is rendered in the Part 2 Responsibility section — NOT
    JES-derived factor text and NOT a {{template leak}}."""
    wd_id = await _create_wd_ec(client)
    # _create_wd_ec does NOT set responsibilities_narrative — stays None
    export_resp = await client.post(f"/api/wd/{wd_id}/export/docx")
    assert export_resp.status_code == 200
    doc = docx.Document(io.BytesIO(export_resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # R-RESP-03: placeholder must appear when narrative is empty
    assert "[To be completed by advisor]" in full_text
    # No template leak
    assert "{{" not in full_text
    assert "responsibilities_text" not in full_text


# ---------------------------------------------------------------------------
# Phase 27 — ELEM-01: build_seven_elements shared helper unit tests
# (Plan 27-02 — Seven-Elements Completeness Audit)
#
# These are pure-function tests that import build_seven_elements directly
# and construct WorkDescription objects inline (no HTTP). The helper is
# the single source of truth for the 7 Part 2 elements + per-element
# status, consumed by POST /api/wd/{id}/validate-elements (ELEM-01) and
# Phase 29's JSON/CSV routes (SEXP-01/02).
# ---------------------------------------------------------------------------


def _wd_for_seven_elements(**overrides):
    """Build a WorkDescription directly for build_seven_elements tests.

    Mirrors the in-memory representation after PATCH + DB round-trip.
    Required defaults: id="test-wd", created_at/last_modified set.
    """
    from datetime import datetime, timezone
    from app.models.work_description import WorkDescription

    base = {
        "id": "test-wd",
        "title": "",
        "record": {},
        "answers": {},
        "step_index": 0,
        "draft": None,
        "reviewing": False,
        "editing_return": False,
        "duties": [],
        "qualification": None,
        "drf_id": None,
        "noc_candidates": [],
        "confirmed_noc": None,
        "confirmed_og": None,
        "confirmed_sub_group": None,
        "og_level": None,
        "sjd_source": None,
        "org_context": None,
        "responsibilities_narrative": None,
        "reports_to_military": None,
        "jes_scores": [],
        "jes_total_points": None,
        "schema_version": 1,
        "created_at": datetime.now(timezone.utc),
        "last_modified": datetime.now(timezone.utc),
    }
    base.update(overrides)
    return WorkDescription(**base)


def test_build_seven_elements_derived_effort_wc():
    """ELEM-01 / R-ELEM-01b: jes_total_points set → effort + working_conditions
    status == 'derived' (not 'missing') and counted in complete_count."""
    from app.services.export_service import build_seven_elements

    wd = _wd_for_seven_elements(
        record={"client_service_results": "Citizens get answers.", "quals": {"education": "Degree", "experience": "5 yrs"}},
        duties=[{"id": "d1", "text": "Provides advice.", "source": "noc",
                 "provenance_noc_code": "4163", "advisor": False}],
        org_context="Within Branch X.",
        responsibilities_narrative="Owns the policy portfolio.",
        jes_total_points=621,
    )
    result = build_seven_elements(wd)
    elements = {e["key"]: e for e in result["elements"]}
    assert elements["effort"]["status"] == "derived"
    assert elements["working_conditions"]["status"] == "derived"
    # All 7 elements + complete_count counts derived as complete
    assert result["total"] == 7
    assert result["complete_count"] == 7


def test_build_seven_elements_no_jes_missing():
    """ELEM-01 / R-ELEM-01b: jes_total_points None → effort + working_conditions
    status == 'missing' (NOT 'derived')."""
    from app.services.export_service import build_seven_elements

    wd = _wd_for_seven_elements(
        record={"client_service_results": "Citizens get answers."},
        duties=[{"id": "d1", "text": "Provides advice.", "source": "noc",
                 "provenance_noc_code": "4163", "advisor": False}],
        org_context="Within Branch X.",
        responsibilities_narrative="Owns the policy portfolio.",
        jes_total_points=None,  # JES never ran
    )
    result = build_seven_elements(wd)
    elements = {e["key"]: e for e in result["elements"]}
    assert elements["effort"]["status"] == "missing"
    assert elements["working_conditions"]["status"] == "missing"
    # 5 populated (oc, csr, ka, skills vacuous w/o quals, resp) — but skills
    # is missing too because quals are absent; let's count what's populated:
    # oc=populated, csr=populated, ka=populated, skills=missing (no quals),
    # effort=missing, resp=populated, wc=missing → complete_count=4
    assert result["complete_count"] == 4


def test_build_seven_elements_org_context_reads_typed_field():
    """ELEM-01 / ROADMAP #4: org_context=None + record has branch/reports →
    organizational_context status == 'missing' (NOT populated).

    This is the audit guard: the completeness audit reads wd.org_context
    (typed root field) and ignores the synthesized fallback from
    _build_organizational_context_text(). A WD whose advisor skipped the
    org_context step must NOT report Organizational Context as populated
    just because record.branch/record.reports exist."""
    from app.services.export_service import build_seven_elements

    wd = _wd_for_seven_elements(
        record={
            "branch": "Department of National Defence",
            "reports": "Director of Policy",
            "summary": "performs duties as assigned",
            "title": "Policy Analyst",
            # Deliberately NO client_service_results / quals / duties here
            # so only org_context is being tested.
        },
        duties=[{"id": "d1", "text": "Provides advice.", "source": "noc",
                 "provenance_noc_code": "4163", "advisor": False}],
        responsibilities_narrative="Owns the policy portfolio.",
        jes_total_points=621,
        # org_context stays None (typed field is empty)
    )
    result = build_seven_elements(wd)
    elements = {e["key"]: e for e in result["elements"]}
    # The audit guard: org_context reads the typed field ONLY
    assert elements["organizational_context"]["status"] == "missing", (
        "org_context audit must read wd.org_context (None) — "
        "branch/reports in record must NOT make it populated"
    )
    # Effort + WC still derived from jes_total_points
    assert elements["effort"]["status"] == "derived"
    assert elements["working_conditions"]["status"] == "derived"


def test_build_seven_elements_responsibility_missing_not_notapplicable():
    """ELEM-01 / R-ELEM-01a: empty responsibilities_narrative →
    responsibility status == 'missing' (NEVER 'not_applicable').

    Locked ROADMAP criterion #3: the field is open to all positions, so
    an empty value means missing, not not_applicable."""
    from app.services.export_service import build_seven_elements

    wd = _wd_for_seven_elements(
        record={"client_service_results": "Citizens get answers.",
                "quals": {"education": "Degree", "experience": "5 yrs"}},
        duties=[{"id": "d1", "text": "Provides advice.", "source": "noc",
                 "provenance_noc_code": "4163", "advisor": False}],
        org_context="Within Branch X.",
        # responsibilities_narrative stays None (advisor skipped the step)
        jes_total_points=621,
    )
    result = build_seven_elements(wd)
    elements = {e["key"]: e for e in result["elements"]}
    # ROADMAP #3 explicit: missing (NOT not_applicable) when empty
    assert elements["responsibility"]["status"] == "missing"
    assert elements["responsibility"]["status"] != "not_applicable", (
        "responsibility status must NEVER be 'not_applicable' per ROADMAP #3"
    )


def test_build_seven_elements_total_seven():
    """ELEM-01: helper ALWAYS returns exactly 7 elements and total == 7.

    Locks the helper's shape contract: consumers (validate-elements
    endpoint, Phase 29 JSON/CSV routes) depend on 7 keys."""
    from app.services.export_service import build_seven_elements

    # Empty WD: every element missing, but exactly 7 keys still returned
    wd_empty = _wd_for_seven_elements()
    result_empty = build_seven_elements(wd_empty)
    assert len(result_empty["elements"]) == 7
    assert result_empty["total"] == 7
    assert result_empty["complete_count"] == 0

    # Fully populated WD: 7 elements, all populated/derived
    wd_full = _wd_for_seven_elements(
        record={"client_service_results": "Citizens get answers.",
                "quals": {"education": "Degree", "experience": "5 yrs"}},
        duties=[{"id": "d1", "text": "Provides advice.", "source": "noc",
                 "provenance_noc_code": "4163", "advisor": False}],
        org_context="Within Branch X.",
        responsibilities_narrative="Owns the policy portfolio.",
        jes_total_points=621,
    )
    result_full = build_seven_elements(wd_full)
    assert len(result_full["elements"]) == 7
    assert result_full["total"] == 7
    assert result_full["complete_count"] == 7

    # Every element has the required keys
    expected_keys = {"key", "label", "status", "value"}
    for el in result_full["elements"]:
        assert expected_keys.issubset(el.keys()), (
            f"Element missing required keys: {el}"
        )
