"""
v2/backend/scripts/build_wd_template.py — Generate the docxtpl TBS Work Description template.

Run from the repo root to (re)generate v2/backend/app/templates/wd_template.docx:

    cd /home/charles/job_description_builder
    python v2/backend/scripts/build_wd_template.py

The .docx is a committed binary artifact — the contract that the export service
(Plan 20-02) fills via a context dict. The Jinja2 variable names listed below
are stable; renaming them is a contract break.

Template structure (TBS Work Description format, adapted from v1.0):
    Title:          "Work Description"
    Section 1:      Position Identification (table: label/value rows)
    Section 2:      Organizational Context (paragraph + italic source citation)
    Section 3:      Summary of Duties — paragraph-level {%p for %} loop
                    with {%p if duty.is_advisor %} gate for advisor-added notes
    Section 4:      Classification & Evaluation — table-row {%tr for %} loop
                    (factor name, degree, points, total)
    Section 5:      Essential Qualifications (education + experience paragraphs)
    Section 6:      Version Manifest — table-row {%tr for %} loop
    Appendix:       Manager Amendments for Review — gated on
                    {%p if amendments|length > 0 %}, paragraph-level loop
                    (AMEND-02)

Jinja2 variables (contract):
    position_title, position_number, og_level, supervisor_title,
    supervisor_position_number, review_date,
    organizational_context_text, organizational_context_source,
    duties (list of {text, noc_code, is_advisor}),
    jes_scores (list of {factor_name, degree, points}),
    jes_total_points,
    education_text, experience_text,
    manifest (list of {source_type, source_id, source_version, retrieved_date}),
    amendments (list of {section, comment, created_at})

Re-run this script to update the template; then commit the regenerated .docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

OUTPUT_PATH = "v2/backend/app/templates/wd_template.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Replace the cell's content with a single paragraph containing the given text."""
    # Clear any default empty paragraph that python-docx places in the cell
    for para in list(cell.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)
    para = cell.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def build() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    title = doc.add_heading("Work Description", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Section 1: Position Identification
    # ------------------------------------------------------------------
    doc.add_heading("1. Position Identification", level=1)

    position_rows = [
        ("Position Title:", "{{ position_title }}"),
        ("Position Number:", "{{ position_number }}"),
        ("Group and Level:", "{{ og_level }}"),
        ("Supervisor Title:", "{{ supervisor_title }}"),
        ("Supervisor Position Number:", "{{ supervisor_position_number }}"),
        ("Review Date:", "{{ review_date }}"),
    ]
    pos_table = doc.add_table(rows=len(position_rows), cols=2)
    pos_table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(position_rows):
        _set_cell_text(pos_table.rows[i].cells[0], label, bold=True)
        _set_cell_text(pos_table.rows[i].cells[1], value)

    # ------------------------------------------------------------------
    # Section 2: Organizational Context
    # ------------------------------------------------------------------
    doc.add_heading("2. Organizational Context", level=1)
    doc.add_paragraph("{{ organizational_context_text }}")
    src_para = doc.add_paragraph()
    src_run = src_para.add_run("Source: {{ organizational_context_source }}")
    src_run.italic = True

    # ------------------------------------------------------------------
    # Section 3: Summary of Duties — paragraph-level loop with advisor gate
    # ------------------------------------------------------------------
    doc.add_heading("3. Summary of Duties", level=1)

    # Each {%p ... %} / {%p endfor %} must be in its own paragraph (per docxtpl:
    # "Do not use {%p twice in the same paragraph").
    doc.add_paragraph("{%p for duty in duties %}")
    doc.add_paragraph("{{ duty.text }}")
    src_para = doc.add_paragraph()
    src_run = src_para.add_run("Source: NOC {{ duty.noc_code }}")
    src_run.italic = True
    # Advisor marker: conditional paragraph with literal label.
    # if/endif each in their own paragraph; the body line between them is
    # rendered only when duty.is_advisor is True.
    doc.add_paragraph("{%p if duty.is_advisor %}")
    doc.add_paragraph("[advisor-added / not from authoritative source]")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endfor %}")

    # ------------------------------------------------------------------
    # Section 4: Classification & Evaluation — JES Scoring table loop
    # ------------------------------------------------------------------
    doc.add_heading("4. Classification & Evaluation", level=1)

    # 4 rows: header, for-marker, data, endfor-marker.
    # The for and endfor MUST be in their own rows (each with only the tag in
    # the first cell) — docxtpl's patch_xml regex is greedy and matches the
    # LAST {%tr %} tag in a row, so co-locating {%tr for %} and {%tr endfor %}
    # in the same data row would cause the for tag to be eaten.
    jes_table = doc.add_table(rows=4, cols=3)
    jes_table.style = "Light Grid Accent 1"
    _set_cell_text(jes_table.rows[0].cells[0], "Factor", bold=True)
    _set_cell_text(jes_table.rows[0].cells[1], "Degree", bold=True)
    _set_cell_text(jes_table.rows[0].cells[2], "Points", bold=True)
    # Row 1: {%tr for %} marker
    _set_cell_text(jes_table.rows[1].cells[0], "{%tr for f in jes_scores %}")
    # Row 2: data row (duplicated per factor by the for loop)
    _set_cell_text(jes_table.rows[2].cells[0], "{{ f.factor_name }}")
    _set_cell_text(jes_table.rows[2].cells[1], "{{ f.degree }}")
    _set_cell_text(jes_table.rows[2].cells[2], "{{ f.points }}")
    # Row 3: {%tr endfor %} marker
    _set_cell_text(jes_table.rows[3].cells[0], "{%tr endfor %}")

    # Total Points line
    total_para = doc.add_paragraph()
    total_run = total_para.add_run("Total Points: ")
    total_run.bold = True
    total_para.add_run("{{ jes_total_points }}")

    # ------------------------------------------------------------------
    # Section 5: Essential Qualifications
    # ------------------------------------------------------------------
    doc.add_heading("5. Essential Qualifications", level=1)

    edu_head = doc.add_paragraph()
    edu_run = edu_head.add_run("Education:")
    edu_run.bold = True
    doc.add_paragraph("{{ education_text }}")

    exp_head = doc.add_paragraph()
    exp_run = exp_head.add_run("Experience:")
    exp_run.bold = True
    doc.add_paragraph("{{ experience_text }}")

    # ------------------------------------------------------------------
    # Section 6: Version Manifest — table-row loop
    # ------------------------------------------------------------------
    doc.add_heading("6. Source Document Version Manifest", level=1)

    # Same for/data/endfor pattern as the JES table.
    manifest_table = doc.add_table(rows=4, cols=4)
    manifest_table.style = "Light Grid Accent 1"
    _set_cell_text(manifest_table.rows[0].cells[0], "Source", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[1], "ID", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[2], "Version", bold=True)
    _set_cell_text(manifest_table.rows[0].cells[3], "Date", bold=True)
    # Row 1: {%tr for %} marker
    _set_cell_text(manifest_table.rows[1].cells[0], "{%tr for m in manifest %}")
    # Row 2: data row
    _set_cell_text(manifest_table.rows[2].cells[0], "{{ m.source_type }}")
    _set_cell_text(manifest_table.rows[2].cells[1], "{{ m.source_id }}")
    _set_cell_text(manifest_table.rows[2].cells[2], "{{ m.source_version }}")
    _set_cell_text(manifest_table.rows[2].cells[3], "{{ m.retrieved_date }}")
    # Row 3: {%tr endfor %} marker
    _set_cell_text(manifest_table.rows[3].cells[0], "{%tr endfor %}")

    # ------------------------------------------------------------------
    # Appendix: Manager Amendments for Review (Phase 20, AMEND-02)
    # Gated on amendments|length > 0 — the advisor may export before any
    # amendment notes are recorded, so the appendix is suppressed in that
    # case to avoid an empty section. The {%p for %} loop iterates the
    # amendments list supplied by export_service.
    # ------------------------------------------------------------------
    doc.add_paragraph("{%p if amendments|length > 0 %}")
    doc.add_heading("Appendix: Manager Amendments for Review", level=1)
    doc.add_paragraph(
        "Manager-proposed amendments — pending advisor ratification."
    )
    # 4 rows: header, for-marker, data, endfor-marker.
    amend_table = doc.add_table(rows=4, cols=3)
    amend_table.style = "Light Grid Accent 1"
    _set_cell_text(amend_table.rows[0].cells[0], "Section", bold=True)
    _set_cell_text(amend_table.rows[0].cells[1], "Comment", bold=True)
    _set_cell_text(amend_table.rows[0].cells[2], "Date", bold=True)
    # Row 1: {%tr for %} marker
    _set_cell_text(amend_table.rows[1].cells[0], "{%tr for a in amendments %}")
    # Row 2: data row
    _set_cell_text(amend_table.rows[2].cells[0], "{{ a.section }}")
    _set_cell_text(amend_table.rows[2].cells[1], "{{ a.comment }}")
    _set_cell_text(amend_table.rows[2].cells[2], "{{ a.created_at }}")
    # Row 3: {%tr endfor %} marker
    _set_cell_text(amend_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()

    # Self-verify: load the generated template with docxtpl and list variables.
    # This catches malformed template tags at build time, not at first export.
    tpl = DocxTemplate(OUTPUT_PATH)
    undeclared = sorted(tpl.get_undeclared_template_variables())
    print(f"WD template variables ({len(undeclared)}): {undeclared}")

    # Contract assertion: all required variables MUST be present in the
    # template's declared variables. Catches a missing section at build time,
    # not at first export.
    required = {
        "position_title", "position_number", "og_level", "supervisor_title",
        "supervisor_position_number", "review_date",
        "organizational_context_text", "organizational_context_source",
        "duties", "jes_scores", "jes_total_points",
        "education_text", "experience_text",
        "manifest", "amendments",
    }
    missing = required - set(undeclared)
    if missing:
        raise AssertionError(
            f"build_wd_template: required variables {missing!r} not declared "
            f"in template. Found: {undeclared}"
        )
    print(f"WD contract: {sorted(required)} declared \u2713")
    print("WD template OK")
