"""
v2/backend/scripts/build_accessible_template.py — Generate the docxtpl GoC Accessible Job Description template.

Run from the repo root to (re)generate v2/backend/app/templates/wd_accessible_template.docx:

    cd /home/charles/job_description_builder
    python v2/backend/scripts/build_accessible_template.py

The .docx is a committed binary artifact — the contract that the export service
(Plan 25-03) fills via a context dict. The Jinja2 variable names listed below
are stable; renaming them is a contract break.

Template structure (GoC Accessible Job Description format, derived from
data/AI Docs/Accessible Job Description Template (1).docx):
    Part 1:         Heading 1 — "Part 1: Position information and signatures"
    Position table: 17-row label/value table (Light Grid Accent 1) —
                    Position number, Position title, Position classification,
                    Position Effective date, Job Code, National occupational
                    classification, Department/Agency Name, Geographic location,
                    Organizational component (Branch/Division), Office code,
                    Language requirements, Linguistic profile, Communications
                    requirements, Security requirements, Supervisor position
                    number, Supervisor position title, Supervisor classification
    Signatures:     3 Heading 2 blocks — Employee statement, Supervisor
                    statement, Manager authorization. Each block: 3 plain
                    paragraphs "Name: ___", "Signature: ___", "Date: ___".
                    NO Jinja2 variables (print-and-sign; locked decision 3).
    Part 2:         Heading 1 — "Part 2: Job description"
    Part 2 subs:    7 Heading 2 subsections — Organizational context, Client
                    service results, Key activities, Skills, Effort,
                    Responsibilities, Working conditions.
    Effort:         table-row {%tr for %} loop over effort_factors; fallback
                    paragraph `{{ effort_placeholder }}` when group has no
                    Effort-category factors (renders "[To be completed by
                    advisor]" or empty string per Plan 25-03 logic).
    Working cond:   table-row {%tr for %} loop over working_conditions_factors;
                    fallback paragraph `{{ wc_placeholder }}` (same pattern).
    Version manifest: {%p for entry in manifest %} paragraph loop, copied from
                    build_wd_template.py verbatim (Plan 25-03 binds manifest
                    data via the same _build_v2_manifest helper).
    Amendments:     {%p if amendments|length > 0 %} gate, {%p for a in
                    amendments %} loop, {%p endif %} — copied verbatim from
                    build_wd_template.py (AMEND-02 contract, Phase 20).

Jinja2 variables (contract):
    Part 1 (position table, 17 fields):
        position_number, position_title, og_level, review_date,
        job_code, noc_code, department_name, geographic_location,
        org_component, office_code, language_requirements, linguistic_profile,
        communications_requirements, security_requirements,
        supervisor_position_number, supervisor_title, supervisor_classification
    Part 2 (7 subsections):
        organizational_context_text, client_service_results_text,
        duties (list of {text, noc_code, is_advisor}),
        education_text, experience_text,
        effort_factors (list of {factor_name, degree, points}),
        effort_placeholder (string — empty when factors present,
            "[To be completed by advisor]" when none),
        responsibilities_text,
        working_conditions_factors (list of {factor_name, degree, points}),
        wc_placeholder (string — same pattern as effort_placeholder)
    Version manifest + amendments:
        manifest (list of {source_type, source_id, source_version, retrieved_date}),
        amendments (list of {section, comment, created_at})

Signature blocks: NO Jinja2 variables (locked decision 3 — print-and-sign only).

Re-run this script to update the template; then commit the regenerated .docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

OUTPUT_PATH = "v2/backend/app/templates/wd_accessible_template.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Replace the cell's content with a single paragraph containing the given text."""
    # Clear any default empty paragraph that python-docx places in the cell
    for para in list(cell.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)
    para = cell.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def build() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ------------------------------------------------------------------
    # PART 1 — Position information and signatures
    # ------------------------------------------------------------------
    part1 = doc.add_heading("Part 1: Position information and signatures", level=1)

    # Position-identification table — 17 rows, label/value, exact order from
    # the reference document. Per locked decision 2, all 17 fields render with
    # a Jinja2 var; the 6 with no current WD data source (Job Code, Office code,
    # Language/Linguistic/Communications/Security requirements) are bound to a
    # [To be completed by advisor] fallback computed in _build_wd_context
    # (Plan 25-03).
    position_rows = [
        ("Position number", "{{ position_number }}"),
        ("Position title", "{{ position_title }}"),
        ("Position classification", "{{ og_level }}"),
        ("Position Effective date", "{{ review_date }}"),
        ("Job Code", "{{ job_code }}"),
        ("National occupational classification", "{{ noc_code }}"),
        ("Department/Agency Name", "{{ department_name }}"),
        ("Geographic location", "{{ geographic_location }}"),
        ("Organizational component (Branch/Division)", "{{ org_component }}"),
        ("Office code", "{{ office_code }}"),
        ("Language requirements", "{{ language_requirements }}"),
        ("Linguistic profile", "{{ linguistic_profile }}"),
        ("Communications requirements", "{{ communications_requirements }}"),
        ("Security requirements", "{{ security_requirements }}"),
        ("Supervisor position number", "{{ supervisor_position_number }}"),
        ("Supervisor position title", "{{ supervisor_title }}"),
        ("Supervisor classification", "{{ supervisor_classification }}"),
    ]
    pos_table = doc.add_table(rows=len(position_rows), cols=2)
    pos_table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(position_rows):
        _set_cell_text(pos_table.rows[i].cells[0], label, bold=True)
        _set_cell_text(pos_table.rows[i].cells[1], value)

    # ------------------------------------------------------------------
    # Signature blocks — 3 Heading 2 sections, static text only (no Jinja2
    # variables). Per locked decision 3, the advisor prints and signs these
    # by hand; they are NOT data-bound to WorkDescription fields.
    # ------------------------------------------------------------------
    for signature_heading in (
        "Employee statement",
        "Supervisor statement",
        "Manager authorization",
    ):
        doc.add_heading(signature_heading, level=2)
        # 3 plain paragraphs (no formatting, no Jinja2) — name/signature/date
        # lines with literal underscore lines for the printed signature.
        doc.add_paragraph("Name: ____________________")
        doc.add_paragraph("Signature: ____________________")
        doc.add_paragraph("Date: ____________________")

    # ------------------------------------------------------------------
    # PART 2 — Job description (7 Heading 2 subsections)
    # ------------------------------------------------------------------
    doc.add_heading("Part 2: Job description", level=1)

    # 1. Organizational context — free text
    doc.add_heading("Organizational context", level=2)
    doc.add_paragraph("{{ organizational_context_text }}")

    # 2. Client service results — free text (Plan 25-03 binds this from
    # record.client_service_results, captured by the conversation flow)
    doc.add_heading("Client service results", level=2)
    doc.add_paragraph("{{ client_service_results_text }}")

    # 3. Key activities — paragraph-level {%p for %} loop over duties, with
    # advisor-gate conditional and NOC source citation. Copied from
    # build_wd_template.py lines 110-121 verbatim — same pattern, same
    # attribute names on the duty dict.
    doc.add_heading("Key activities", level=2)
    doc.add_paragraph("{%p for duty in duties %}")
    doc.add_paragraph("{{ duty.text }}")
    src_para = doc.add_paragraph()
    src_run = src_para.add_run("Source: NOC {{ duty.noc_code }}")
    src_run.italic = True
    # Advisor marker: conditional paragraph with literal label.
    doc.add_paragraph("{%p if duty.is_advisor %}")
    doc.add_paragraph("[advisor-added / not from authoritative source]")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endfor %}")

    # 4. Skills — Education + Experience (mirrors the TBS template's Section 5
    # mapping, per PATTERNS.md assumption A1)
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Education: {{ education_text }}")
    doc.add_paragraph("Experience: {{ experience_text }}")

    # 5. Effort — table-row {%tr for %} loop over effort_factors; fallback
    # paragraph below the table renders effort_placeholder (which Plan 25-03
    # sets to "[To be completed by advisor]" when the OG has no Effort-category
    # factors, or "" when factors are present).
    #
    # PITFALL (preserve verbatim from build_wd_template.py): docxtpl's
    # patch_xml regex is greedy and matches the LAST {%tr %} tag in a row —
    # never co-locate for/endfor, never share a row with data.
    doc.add_heading("Effort", level=2)
    effort_table = doc.add_table(rows=4, cols=3)
    effort_table.style = "Light Grid Accent 1"
    _set_cell_text(effort_table.rows[0].cells[0], "Factor", bold=True)
    _set_cell_text(effort_table.rows[0].cells[1], "Degree", bold=True)
    _set_cell_text(effort_table.rows[0].cells[2], "Points", bold=True)
    # Row 1: {%tr for %} marker — MUST be alone in its row, first cell only
    _set_cell_text(effort_table.rows[1].cells[0], "{%tr for f in effort_factors %}")
    # Row 2: data row (duplicated per factor by the for loop)
    _set_cell_text(effort_table.rows[2].cells[0], "{{ f.factor_name }}")
    _set_cell_text(effort_table.rows[2].cells[1], "{{ f.degree }}")
    _set_cell_text(effort_table.rows[2].cells[2], "{{ f.points }}")
    # Row 3: {%tr endfor %} marker — MUST be alone in its row
    _set_cell_text(effort_table.rows[3].cells[0], "{%tr endfor %}")
    # Fallback placeholder — always rendered; Plan 25-03 sets the variable
    # to "" when factors are present and to "[To be completed by advisor]"
    # when the OG has no Effort-category factors at all.
    doc.add_paragraph("{{ effort_placeholder }}")

    # 6. Responsibilities — free text (Plan 25-03 binds this from the
    # Responsibility-category JES factor rationale, falling back to
    # [To be completed by advisor] for groups with no Responsibility factors)
    doc.add_heading("Responsibilities", level=2)
    doc.add_paragraph("{{ responsibilities_text }}")

    # 7. Working conditions — same {%tr for %} structure as Effort, with
    # wc_placeholder fallback. Same pitfall applies (for/endfor alone in
    # their own rows).
    doc.add_heading("Working conditions", level=2)
    wc_table = doc.add_table(rows=4, cols=3)
    wc_table.style = "Light Grid Accent 1"
    _set_cell_text(wc_table.rows[0].cells[0], "Factor", bold=True)
    _set_cell_text(wc_table.rows[0].cells[1], "Degree", bold=True)
    _set_cell_text(wc_table.rows[0].cells[2], "Points", bold=True)
    _set_cell_text(wc_table.rows[1].cells[0], "{%tr for f in working_conditions_factors %}")
    _set_cell_text(wc_table.rows[2].cells[0], "{{ f.factor_name }}")
    _set_cell_text(wc_table.rows[2].cells[1], "{{ f.degree }}")
    _set_cell_text(wc_table.rows[2].cells[2], "{{ f.points }}")
    _set_cell_text(wc_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{{ wc_placeholder }}")

    # ------------------------------------------------------------------
    # Version manifest — paragraph-level {%p for %} loop. Copied verbatim
    # from build_wd_template.py. Plan 25-03 binds manifest via the same
    # _build_v2_manifest helper that populates the TBS template today.
    # ------------------------------------------------------------------
    doc.add_heading("Source Document Version Manifest", level=1)
    doc.add_paragraph("{%p for entry in manifest %}")
    doc.add_paragraph("{{ entry.source_type }} — {{ entry.source_id }} (v{{ entry.source_version }}, retrieved {{ entry.retrieved_date }})")
    doc.add_paragraph("{%p endfor %}")

    # ------------------------------------------------------------------
    # Appendix: Manager Amendments for Review (AMEND-02, gated on
    # amendments|length > 0). Copied verbatim from build_wd_template.py —
    # the export_service already populates amendments via _get_amendments
    # and Plan 25-03 preserves that contract.
    # ------------------------------------------------------------------
    doc.add_paragraph("{%p if amendments|length > 0 %}")
    doc.add_heading("Appendix: Manager Amendments for Review", level=1)
    doc.add_paragraph(
        "Manager-proposed amendments — pending advisor ratification."
    )
    amend_table = doc.add_table(rows=4, cols=3)
    amend_table.style = "Light Grid Accent 1"
    _set_cell_text(amend_table.rows[0].cells[0], "Section", bold=True)
    _set_cell_text(amend_table.rows[0].cells[1], "Comment", bold=True)
    _set_cell_text(amend_table.rows[0].cells[2], "Date", bold=True)
    _set_cell_text(amend_table.rows[1].cells[0], "{%tr for a in amendments %}")
    _set_cell_text(amend_table.rows[2].cells[0], "{{ a.section }}")
    _set_cell_text(amend_table.rows[2].cells[1], "{{ a.comment }}")
    _set_cell_text(amend_table.rows[2].cells[2], "{{ a.created_at }}")
    _set_cell_text(amend_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()

    # Self-verify: load the generated template with docxtpl and list variables.
    # This catches malformed template tags at build time, not at first export.
    tpl = DocxTemplate(OUTPUT_PATH)
    undeclared = sorted(tpl.get_undeclared_template_variables())
    print(f"Accessible template variables ({len(undeclared)}): {undeclared}")

    # Contract assertion: all required variables MUST be present in the
    # template's declared variables. Catches a missing section at build time,
    # not at first export.
    required = {
        "position_number", "position_title", "og_level", "review_date",
        "job_code", "noc_code", "department_name", "geographic_location",
        "org_component", "office_code", "language_requirements", "linguistic_profile",
        "communications_requirements", "security_requirements",
        "supervisor_position_number", "supervisor_title", "supervisor_classification",
        "organizational_context_text", "client_service_results_text",
        "duties", "education_text", "experience_text",
        "effort_factors", "effort_placeholder",
        "responsibilities_text",
        "working_conditions_factors", "wc_placeholder",
        "manifest", "amendments",
    }
    missing = required - set(undeclared)
    if missing:
        raise AssertionError(
            f"build_accessible_template: required variables {missing!r} not declared "
            f"in template. Found: {undeclared}"
        )
    print(f"Accessible contract: {sorted(required)} declared \u2713")
    print("Accessible template OK")
