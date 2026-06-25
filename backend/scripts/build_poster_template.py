"""
backend/scripts/build_poster_template.py — Generate the docxtpl Job Poster template.

Run from the repo root to (re)generate backend/app/templates/poster_template.docx:

    cd /home/charles/job_description_builder
    python backend/scripts/build_poster_template.py

The .docx is a committed binary artifact — the contract that the export service
(Plan 20-02) fills via a context dict. The Jinja2 variable names listed below
are stable; renaming them is a contract break.

Template structure (Bilingual job poster, EXP-02):
    Header:         "JOB POSTER / AFFICHE D'EMPLOI"
    Position/Poste: {{ position_title }} / {{ bilingual_title_fr }}
    OG & Level:     {{ og_level }} — {{ og_name }}
    Branch:         {{ branch }}
    Key Duties:     paragraph-level {%p for %} loop over duties (top 3-5)
    Qualifications: {{ education_text }} and {{ experience_text }}
    How to Apply:   static placeholder text

Jinja2 variables (contract):
    position_title, bilingual_title_fr,
    og_level, og_name, branch,
    duties (list of {text} — top 3-5 entries),
    education_text, experience_text

Re-run this script to update the template; then commit the regenerated .docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

OUTPUT_PATH = "backend/app/templates/poster_template.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Replace the cell's content with a single paragraph containing the given text."""
    for para in list(cell.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)
    para = cell.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def build() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ------------------------------------------------------------------
    # Bilingual header
    # ------------------------------------------------------------------
    title = doc.add_heading("JOB POSTER / AFFICHE D'EMPLOI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Position / Poste (bilingual)
    # ------------------------------------------------------------------
    pos_head = doc.add_paragraph()
    pos_run = pos_head.add_run("Position / Poste:")
    pos_run.bold = True
    doc.add_paragraph("{{ position_title }}")
    doc.add_paragraph("{{ bilingual_title_fr }}")

    # ------------------------------------------------------------------
    # OG & Level / Groupe et niveau
    # ------------------------------------------------------------------
    og_head = doc.add_paragraph()
    og_run = og_head.add_run("Group and Level / Groupe et niveau:")
    og_run.bold = True
    doc.add_paragraph("{{ og_level }} \u2014 {{ og_name }}")

    # ------------------------------------------------------------------
    # Branch / Direction
    # ------------------------------------------------------------------
    branch_head = doc.add_paragraph()
    branch_run = branch_head.add_run("Branch / Direction:")
    branch_run.bold = True
    doc.add_paragraph("{{ branch }}")

    # ------------------------------------------------------------------
    # About the Organization / À propos de l'organisation
    # ------------------------------------------------------------------
    org_head = doc.add_paragraph()
    org_run = org_head.add_run("About the Organization / À propos de l'organisation:")
    org_run.bold = True
    doc.add_paragraph("{{ org_context }}")

    # ------------------------------------------------------------------
    # Key Duties / Principales fonctions — paragraph-level loop
    # The export service slices the full duties list to top 3-5 entries
    # before passing the context dict to the template.
    # ------------------------------------------------------------------
    doc.add_heading("Key Duties / Principales fonctions", level=1)

    # Each {%p ... %} / {%p endfor %} must be in its own paragraph (per docxtpl).
    doc.add_paragraph("{%p for duty in duties %}")
    doc.add_paragraph("\u2022 {{ duty.text }}")
    doc.add_paragraph("{%p endfor %}")

    # ------------------------------------------------------------------
    # Qualifications
    # ------------------------------------------------------------------
    doc.add_heading("Qualifications", level=1)

    edu_head = doc.add_paragraph()
    edu_run = edu_head.add_run("Education / \u00c9tudes:")
    edu_run.bold = True
    doc.add_paragraph("{{ education_text }}")

    exp_head = doc.add_paragraph()
    exp_run = exp_head.add_run("Experience / Exp\u00e9rience:")
    exp_run.bold = True
    doc.add_paragraph("{{ experience_text }}")

    # ------------------------------------------------------------------
    # How to Apply / Comment postuler — static placeholder
    # ------------------------------------------------------------------
    doc.add_heading("How to Apply / Comment postuler", level=1)
    doc.add_paragraph("[To be provided by HR / \u00c0 fournir par les RH]")

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()

    # Self-verify: load the generated template with docxtpl and list variables.
    # This catches malformed template tags at build time, not at first export.
    tpl = DocxTemplate(OUTPUT_PATH)
    undeclared = sorted(tpl.get_undeclared_template_variables())
    print(f"Poster template variables ({len(undeclared)}): {undeclared}")

    # Contract assertion: all required variables MUST be present in the
    # template's declared variables. Catches a missing section at build time.
    required = {
        "position_title", "bilingual_title_fr",
        "og_level", "og_name", "branch",
        "duties", "education_text", "experience_text",
        "org_context",
    }
    missing = required - set(undeclared)
    if missing:
        raise AssertionError(
            f"build_poster_template: required variables {missing!r} not declared "
            f"in template. Found: {undeclared}"
        )
    print(f"Poster contract: {sorted(required)} declared \u2713")
    print("Poster template OK")
